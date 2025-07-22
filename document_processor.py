import os
import requests
from typing import List, Optional
from urllib.parse import urlparse
import PyPDF2
from docx import Document
from bs4 import BeautifulSoup
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangChainDocument
from config import Config

class DocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=Config.CHUNK_SIZE,
            chunk_overlap=Config.CHUNK_OVERLAP,
            length_function=len,
        )
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            raise Exception(f"Error extracting text from PDF: {str(e)}")
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            raise Exception(f"Error extracting text from DOCX: {str(e)}")
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            raise Exception(f"Error extracting text from TXT: {str(e)}")
    
    def extract_text_from_url(self, url: str) -> str:
        """Extract text from URL."""
        try:
            response = requests.get(url, timeout=30)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text
            text = soup.get_text()
            
            # Clean up text
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            return text
        except Exception as e:
            raise Exception(f"Error extracting text from URL: {str(e)}")
    
    def process_file(self, file_path: str) -> List[LangChainDocument]:
        """Process uploaded file and return chunks."""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            text = self.extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            text = self.extract_text_from_docx(file_path)
        elif file_extension == '.txt':
            text = self.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        return self.create_chunks(text, source=file_path)
    
    def process_url(self, url: str) -> List[LangChainDocument]:
        """Process URL and return chunks."""
        text = self.extract_text_from_url(url)
        return self.create_chunks(text, source=url)
    
    def create_chunks(self, text: str, source: str) -> List[LangChainDocument]:
        """Create text chunks from extracted text."""
        chunks = self.text_splitter.split_text(text)
        
        documents = []
        for i, chunk in enumerate(chunks):
            doc = LangChainDocument(
                page_content=chunk,
                metadata={
                    "source": source,
                    "chunk_id": i,
                    "total_chunks": len(chunks)
                }
            )
            documents.append(doc)
        
        return documents